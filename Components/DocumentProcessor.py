from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import aspose.words as aw


def DisableLinks(filepath):
    print('=================Remove Links===================')
    doc = Document(filepath)
    rels = doc.part.rels
    logs = []
    for rel in rels:
        if rels[rel].reltype == RT.HYPERLINK:
            log = f'[-]Remove the hyperlink ==>{rels[rel]._target}'
            logs.append(log)
            print(log)
            doc.part.rels[rel]._target='#'
    doc.save(filepath)
    print('[=]Successfully remove all Links.')  
    print('==============================================================')
    return logs if logs != None else None 


def DisableMacros(filepath,filename):
    print('================Remove Macros================')
    doc = aw.Document(filepath)
    vba_list =list(doc.vba_project.modules)
    logs =  []
    for _ in range(len(vba_list)):
        log = f'[-] Remove the Macro ==>> {doc.vba_project.modules[0].name}'
        print(log)
        logs.append(log)
        doc.vba_project.modules.remove(doc.vba_project.modules[0])
    
    doc.remove_macros()
    doc.save(filename+'.docx') 
    clean_docx(filename+'.docx')
    print('[=]Successfully remove all Macros.')  
    print('==============================================================')
    return logs if logs != None else None 
       



# need to fix ... 
def clean_docx(filepath):
    clean_footer(filepath)
    clean_paragraphs(filepath) 



first_line = 'Evaluation Only. Created with Aspose.Words. Copyright 2003-2022 Aspose Pty Ltd.'
footer_clean = 'Created with an evaluation copy of Aspose.Words. To discover the full versions of our APIs please visit: https://products.aspose.com/words/'

def clean_footer(filepath):
    doc = Document(filepath)
    for section in doc.sections:
        footer = section.footer
        for i in range(len(footer.paragraphs)):
            footer.paragraphs[i].text = footer.paragraphs[i].text.replace(footer_clean,'')
    doc.save(filepath) 
         

def clean_paragraphs(filepath):
    doc = Document(filepath)
    for i in range(len(doc.paragraphs)):
        doc.paragraphs[i].text = doc.paragraphs[i].text.replace(first_line,'')
    doc.save(filepath)


